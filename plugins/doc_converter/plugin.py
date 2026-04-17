import os
import json
import csv
from typing import List
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox,
)
from core.plugin_base import PluginBase, PluginMeta


class DocConverterPlugin(PluginBase):

    @property
    def meta(self) -> PluginMeta:
        return PluginMeta(
            id="doc_converter",
            name="Doc Converter",
            description="Convert DOCX to TXT/HTML, XLSX to CSV/JSON, and Markdown to HTML.",
            category="Documents",
            icon_char="📝",
            icon_name="documents",
            accent_color="#5BA85B",
            supported_inputs=["docx", "xlsx", "md", "markdown"],
            supported_outputs=["txt", "html", "csv", "json"],
            requires=["python-docx>=1.1.0", "openpyxl>=3.1.0"],
        )

    def build_ui(self, parent: QWidget) -> QWidget:
        w = QWidget(parent)
        layout = QVBoxLayout(w)
        layout.setSpacing(12)

        row = QHBoxLayout()
        row.addWidget(QLabel("Output format:"))
        self._fmt = QComboBox()
        self._fmt.addItems(["Auto (by input type)", "TXT", "HTML", "CSV", "JSON"])
        row.addWidget(self._fmt)
        row.addStretch()
        layout.addLayout(row)

        layout.addWidget(QLabel(
            "Auto:\n"
            "  .docx → TXT\n"
            "  .xlsx → CSV\n"
            "  .md   → HTML"
        ))
        layout.addStretch()
        return w

    def run(self, input_files: List[str], output_dir: str) -> None:
        fmt = self._fmt.currentText()
        total = len(input_files)

        for i, path in enumerate(input_files):
            if self.is_cancelled:
                self.emit_done(False, "Abgebrochen.")
                return

            ext = os.path.splitext(path)[1].lower().lstrip(".")
            try:
                if ext == "docx":
                    self._convert_docx(path, output_dir, fmt)
                elif ext == "xlsx":
                    self._convert_xlsx(path, output_dir, fmt)
                elif ext in ("md", "markdown"):
                    self._convert_md(path, output_dir)
                else:
                    self.emit_log(f"✗  Nicht unterstützt: {ext}", "warn")
            except Exception as e:
                self.emit_log(f"✗  {os.path.basename(path)}: {e}", "error")

            self.emit_progress(int((i + 1) / total * 100))

        self.emit_done(True, f"{total} Datei(en) konvertiert.")

    def _convert_docx(self, path: str, output_dir: str, fmt: str):
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx nicht installiert.")

        doc = docx.Document(path)
        base = os.path.splitext(os.path.basename(path))[0]
        target_fmt = "TXT" if fmt == "Auto (by input type)" else fmt

        if target_fmt in ("TXT", "Auto (by input type)"):
            text = "\n".join(p.text for p in doc.paragraphs)
            out = os.path.join(output_dir, f"{base}.txt")
            with open(out, "w", encoding="utf-8") as f:
                f.write(text)
        else:
            lines = [f"<p>{p.text}</p>" for p in doc.paragraphs if p.text]
            html = "<html><body>\n" + "\n".join(lines) + "\n</body></html>"
            out = os.path.join(output_dir, f"{base}.html")
            with open(out, "w", encoding="utf-8") as f:
                f.write(html)

        self.emit_log(f"✓  {os.path.basename(path)} → {os.path.basename(out)}", "ok")

    def _convert_xlsx(self, path: str, output_dir: str, fmt: str):
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl nicht installiert.")

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        base = os.path.splitext(os.path.basename(path))[0]
        target_fmt = "CSV" if fmt == "Auto (by input type)" else fmt

        for sheet in wb.sheetnames:
            ws = wb[sheet]
            rows = [[str(cell.value or "") for cell in row] for row in ws.iter_rows()]
            suffix = f"_{sheet}" if len(wb.sheetnames) > 1 else ""

            if target_fmt == "JSON":
                if rows:
                    headers, *data = rows
                    result = [dict(zip(headers, row)) for row in data]
                else:
                    result = []
                out = os.path.join(output_dir, f"{base}{suffix}.json")
                with open(out, "w", encoding="utf-8") as f:
                    json.dump(result, f, ensure_ascii=False, indent=2)
            else:
                out = os.path.join(output_dir, f"{base}{suffix}.csv")
                with open(out, "w", newline="", encoding="utf-8") as f:
                    csv.writer(f).writerows(rows)

            self.emit_log(f"✓  {os.path.basename(path)}[{sheet}] → {os.path.basename(out)}", "ok")

    def _convert_md(self, path: str, output_dir: str):
        base = os.path.splitext(os.path.basename(path))[0]
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()

        try:
            import markdown
            html_body = markdown.markdown(text)
        except ImportError:
            # Simple fallback ohne markdown-Paket
            html_body = "<pre>" + text.replace("&", "&amp;").replace("<", "&lt;") + "</pre>"

        html = f"<html><body>\n{html_body}\n</body></html>"
        out = os.path.join(output_dir, f"{base}.html")
        with open(out, "w", encoding="utf-8") as f:
            f.write(html)
        self.emit_log(f"✓  {os.path.basename(path)} → {os.path.basename(out)}", "ok")
